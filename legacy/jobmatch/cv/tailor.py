"""Adaptación del CV a una vacante concreta (tailoring) con python-docx.

Cambios respecto al original:
- reutiliza la capa de IA (mismo cliente, reintentos y parseo) en vez de
  crear un segundo cliente de Gemini,
- `apply_cv_strategy` se dividió en pasos y protege contra secciones que no
  se encuentren (antes reventaba con KeyError),
- `extract_skills` ahora reconoce también el CV en español (relevante ahora
  que las búsquedas son en español).
"""

from __future__ import annotations

from docx import Document
from langdetect import detect

from jobmatch.pipeline import ai
from jobmatch.pipeline.prompts import CV_TAILOR_PROMPT

_SKILLS_TITLES = {"CORE SKILLS", "HABILIDADES TÉCNICAS"}
_EXPERIENCE_TITLES = {"PROFESSIONAL EXPERIENCE", "EXPERIENCIA PROFESIONAL"}

_SECTION_TITLES = {
    "summary": {"PROFESSIONAL SUMMARY", "PERFIL PROFESIONAL"},
    "skills": _SKILLS_TITLES,
    "experience": _EXPERIENCE_TITLES,
    "education": {"EDUCATION", "EDUCACIÓN"},
}


def detect_job_language(job_description: str) -> str:
    try:
        return "es" if detect(job_description) == "es" else "en"
    except Exception:
        return "en"


def read_cv_docx(file_path: str) -> str:
    document = Document(file_path)
    return "\n".join(p.text.strip() for p in document.paragraphs if p.text.strip())


def extract_skills(cv_content: str) -> list[str]:
    skills: list[str] = []
    inside = False
    for raw_line in cv_content.split("\n"):
        line = raw_line.strip()
        if line in _SKILLS_TITLES:
            inside = True
            continue
        if line in _EXPERIENCE_TITLES:
            break
        if inside and ":" in line:
            skills.append(line.split(":")[0].strip())
    return skills


def generate_cv_strategy(cv_content: str, job_description: str, skills) -> dict:
    prompt = CV_TAILOR_PROMPT.format(
        cv_content=cv_content,
        job_description=job_description,
        skills=skills,
    )
    strategy = ai.generate_json(prompt)
    if not isinstance(strategy, dict):
        raise ValueError("La IA no devolvió una estrategia de CV válida.")
    return strategy


def find_section_indexes(paragraphs) -> dict:
    indexes: dict[str, int] = {}
    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        for key, names in _SECTION_TITLES.items():
            if text in names:
                indexes[key] = i
    return indexes


def _apply_summary(paragraphs, sections, strategy) -> None:
    if not strategy.get("summary") or "summary" not in sections:
        return
    idx = sections["summary"] + 1
    if idx >= len(paragraphs) or not paragraphs[idx].runs:
        return
    paragraph = paragraphs[idx]
    paragraph.runs[0].text = strategy["summary"]
    for run in paragraph.runs[1:]:
        run.text = ""


def _apply_skills(paragraphs, sections, strategy) -> None:
    order = strategy.get("skills_priority", [])
    if not order or "skills" not in sections or "experience" not in sections:
        return

    start = sections["skills"] + 1
    end = sections["experience"]

    data: dict[str, dict] = {}
    for i in range(start, end):
        paragraph = paragraphs[i]
        if len(paragraph.runs) < 2:
            continue
        category = paragraph.runs[0].text.replace(":", "").strip()
        data[category] = {
            "title": paragraph.runs[0].text,
            "content": paragraph.runs[1].text,
        }

    reordered = [data[category] for category in order if category in data]
    for i, skill in enumerate(reordered, start=start):
        paragraphs[i].runs[0].text = skill["title"]
        paragraphs[i].runs[1].text = skill["content"]


def _apply_bullets(paragraphs, strategy) -> None:
    replacements = {
        item["original"]: item["replacement"]
        for item in strategy.get("bullet_updates", [])
        if "original" in item and "replacement" in item
    }
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if text in replacements and paragraph.runs:
            paragraph.runs[0].text = replacements[text]


def apply_cv_strategy(source_docx: str, strategy: dict, output_docx: str) -> None:
    document = Document(source_docx)
    paragraphs = document.paragraphs
    sections = find_section_indexes(paragraphs)

    _apply_summary(paragraphs, sections, strategy)
    _apply_skills(paragraphs, sections, strategy)
    _apply_bullets(paragraphs, strategy)

    document.save(output_docx)
